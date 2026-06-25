# tests/test_nodes.py
"""Unit tests for the document analysis pipeline nodes."""
from unittest.mock import MagicMock, patch
import pytest
import sys

# ── Mock llm_guard AVANT tout import de nodes ──
sys.modules["llm_guard"] = MagicMock()
sys.modules["llm_guard.vault"] = MagicMock()
sys.modules["llm_guard.input_scanners"] = MagicMock()

# ── Mock HuggingFaceEmbeddings AVANT import de retrieval ──
sys.modules["langchain_huggingface"] = MagicMock()


# ─────────────────────────────────────────────
# EXTRACTION — extract_text
# ─────────────────────────────────────────────


class TestExtractText:
    def test_pdf_extraction(self, tmp_path):
        """extract_text returns text from a valid PDF."""
        import fitz
        from nodes.extraction import extract_text

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Hello PDF")
        doc.save(str(pdf_path))
        doc.close()

        result = extract_text(str(pdf_path))
        assert "Hello PDF" in result

    def test_docx_extraction(self, tmp_path):
        """extract_text returns text from a valid DOCX."""
        from docx import Document as DocxDocument
        from nodes.extraction import extract_text

        docx_path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.add_paragraph("Hello DOCX")
        doc.save(str(docx_path))

        result = extract_text(str(docx_path))
        assert "Hello DOCX" in result

    def test_doc_raises_value_error(self):
        """extract_text raises ValueError for .doc files."""
        from nodes.extraction import extract_text

        with pytest.raises(ValueError, match=".doc not supported"):
            extract_text("file.doc")

    def test_unsupported_format_raises_value_error(self):
        """extract_text raises ValueError for unknown formats."""
        from nodes.extraction import extract_text

        with pytest.raises(ValueError, match="Unsupported format"):
            extract_text("file.txt")

    def test_docx_no_non_breaking_spaces(self, tmp_path):
        """extract_text replaces \\xa0 with regular spaces."""
        from docx import Document as DocxDocument
        from nodes.extraction import extract_text

        docx_path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.add_paragraph("Hello\xa0World")
        doc.save(str(docx_path))

        result = extract_text(str(docx_path))
        assert "\xa0" not in result


# ─────────────────────────────────────────────
# EXTRACTION — extraction_node
# ─────────────────────────────────────────────


class TestExtractionNode:
    def test_empty_input_path_returns_error(self):
        """extraction_node returns error if input_path is empty."""
        from nodes.extraction import extraction_node

        result = extraction_node({"input_path": "", "context": "context"})
        assert "Error" in result["rag_result"]

    def test_unsupported_format_returns_error(self):
        """extraction_node returns error for unsupported format."""
        from nodes.extraction import extraction_node

        result = extraction_node({"input_path": "file.txt", "context": "context"})
        assert "rag_result" in result

    def test_valid_pdf_returns_document_content(self, tmp_path):
        """extraction_node returns document_content for a valid PDF."""
        import fitz
        from nodes.extraction import extraction_node

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Python developer")
        doc.save(str(pdf_path))
        doc.close()

        with patch(
            "nodes.extraction.scan_prompt", return_value=("Python developer", True, 0.0)
        ):
            result = extraction_node(
                {"input_path": str(pdf_path), "context": "Python job"}
            )

        assert "document_content" in result
        assert "context_text" in result

    def test_empty_context_continues(self, tmp_path):
        """extraction_node continues without error if context is empty."""
        import fitz
        from nodes.extraction import extraction_node

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Developer")
        doc.save(str(pdf_path))
        doc.close()

        with patch(
            "nodes.extraction.scan_prompt", return_value=("Developer", True, 0.0)
        ):
            result = extraction_node({"input_path": str(pdf_path), "context": ""})

        assert "document_content" in result
        assert result["context_text"] == ""

    def test_risky_document_returns_error(self, tmp_path):
        """extraction_node blocks if document content is risky."""
        import fitz
        from nodes.extraction import extraction_node

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "some text")
        doc.save(str(pdf_path))
        doc.close()

        with patch("nodes.extraction.scan_prompt", return_value=("", False, 1.0)):
            result = extraction_node(
                {"input_path": str(pdf_path), "context": "context"}
            )

        assert "Error" in result["rag_result"]


# ─────────────────────────────────────────────
# RETRIEVAL
# ─────────────────────────────────────────────


class TestRetrievalNode:
    def test_empty_document_returns_error(self):
        """retrieval_node returns error if document_content is empty."""
        from nodes.retrieval import retrieval_node

        result = retrieval_node({"document_content": "", "context_text": "query"})
        assert "Error" in result["rag_result"]

    def test_empty_context_returns_error(self):
        """retrieval_node returns error if context_text is empty."""
        from nodes.retrieval import retrieval_node

        result = retrieval_node({"document_content": "some doc", "context_text": ""})
        assert "Error" in result["rag_result"]

    def test_empty_scan_marker_returns_error(self):
        """retrieval_node returns error for __EMPTY_SCAN__ marker."""
        from nodes.retrieval import retrieval_node

        result = retrieval_node(
            {"document_content": "__EMPTY_SCAN__", "context_text": "query"}
        )
        assert "Error" in result["rag_result"]

    def test_valid_input_returns_dict(self):
        """retrieval_node returns dict with relevant_docs and query."""
        from langchain_core.documents import Document
        from nodes.retrieval import retrieval_node

        mock_docs = [Document(page_content="Python experience")]
        mock_vectorstore = MagicMock()
        mock_vectorstore.similarity_search.return_value = mock_docs

        with patch("nodes.retrieval.Chroma.from_texts", return_value=mock_vectorstore):
            result = retrieval_node(
                {
                    "document_content": "Python developer with 5 years",
                    "context_text": "Looking for Python developer",
                }
            )

        assert isinstance(result["rag_result"], dict)
        assert "relevant_docs" in result["rag_result"]
        assert "query" in result["rag_result"]
        assert result["rag_result"]["query"] == "Looking for Python developer"

    def test_exception_returns_error(self):
        """retrieval_node captures exceptions and returns error."""
        from nodes.retrieval import retrieval_node

        with patch(
            "nodes.retrieval.Chroma.from_texts", side_effect=Exception("chroma error")
        ):
            result = retrieval_node(
                {"document_content": "some doc", "context_text": "some query"}
            )

        assert "Error" in result["rag_result"]


# ─────────────────────────────────────────────
# ANALYSIS
# ─────────────────────────────────────────────


class TestAnalysisNode:
    def _make_state(self, rag_result=None):
        from langchain_core.documents import Document

        return {
            "input_path": "document.pdf",
            "context": "query",
            "document_content": "some document",
            "context_text": "some query",
            "rag_result": (
                rag_result
                if rag_result is not None
                else {
                    "relevant_docs": [Document(page_content="Python experience")],
                    "query": "Looking for a Python developer",
                }
            ),
            "final_output": None,
        }

    def test_empty_rag_result_returns_empty_report(self):
        """analysis_node returns empty report if rag_result is empty."""
        from nodes.analysis import analysis_node

        result = analysis_node(self._make_state(rag_result={}))
        assert result["final_output"] == {}

    def test_string_rag_result_returns_empty_report(self):
        """analysis_node returns empty report if rag_result is an error string."""
        from nodes.analysis import analysis_node

        result = analysis_node(self._make_state(rag_result="Error: invalid document"))
        assert result["final_output"] == {}

    def test_valid_input_calls_llm(self):
        """analysis_node calls the LLM for valid input."""
        from nodes.analysis import analysis_node
        from models.schemas import AnalysisResult

        mock_response = AnalysisResult(
            score_a=80, score_b=70, suggestions=["Add metrics"]
        )
        mock_chain = MagicMock()
        mock_chain.invoke.return_value = mock_response

        with patch("nodes.analysis.ChatPromptTemplate.from_messages") as mock_prompt:
            mock_prompt.return_value.__or__ = MagicMock(return_value=mock_chain)
            result = analysis_node(self._make_state())

        assert "final_output" in result

    def test_llm_exception_returns_empty_report(self):
        """analysis_node returns empty report if LLM fails."""
        from nodes.analysis import analysis_node

        mock_chain = MagicMock()
        mock_chain.invoke.side_effect = Exception("LLM timeout")

        with patch("nodes.analysis.ChatPromptTemplate.from_messages") as mock_prompt:
            mock_prompt.return_value.__or__ = MagicMock(return_value=mock_chain)
            result = analysis_node(self._make_state())

        assert result["final_output"] == {}
